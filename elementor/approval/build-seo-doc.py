"""Builds the client approval document as a real .docx, tables and all.

Generated from doc-data.json, which is generated from elementor/seo.mjs, for
the same reason the artifact and the email are: the strings a client approves
have to be the strings that deploy, and a hand-typed table is a second copy
that drifts silently.

Landscape, because the description column is the point of the document and a
160-character line in a portrait table wraps into a column of soup.
"""
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).parent
DATA = json.loads((HERE / "doc-data.json").read_text())

NAVY = RGBColor(0x00, 0x3C, 0x50)
ORANGE = RGBColor(0xE6, 0x5A, 0x28)
MUTED = RGBColor(0x5C, 0x6E, 0x77)
INK = RGBColor(0x11, 0x24, 0x2D)

SITE = "empowerms.org"

doc = Document()

# Landscape letter, narrow margins: four columns of real text need the width.
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
for attr in ("left_margin", "right_margin"):
    setattr(section, attr, Inches(0.6))
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)

base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(10)
base.font.color.rgb = INK


def shade(cell, hex_fill):
    """Cell background. python-docx has no API for it; this is the XML."""
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def heading(text, size=18, color=NAVY, space_before=18, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def body(text, size=10, color=INK, italic=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    return p


def rule():
    """The site's own motif: a short orange bar above a block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("▆▆▆")
    run.font.size = Pt(8)
    run.font.color.rgb = ORANGE
    return p


def build_table(rows, show_today):
    cols = ["Page", "Proposed title", "Proposed description", "Today"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(1.9), Inches(2.4), Inches(4.0), Inches(1.5)]

    header = table.rows[0]
    for i, label in enumerate(cols):
        cell = header.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(label.upper())
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, "003C50")

    for r in rows:
        cells = table.add_row().cells

        # Page: readable name, then the address beneath it in small grey.
        p = cells[0].paragraphs[0]
        run = p.add_run(r["name"])
        run.font.bold = True
        run.font.size = Pt(10)
        p2 = cells[0].add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        run2 = p2.add_run(SITE + r["path"])
        run2.font.size = Pt(7.5)
        run2.font.color.rgb = MUTED

        # Title, with its character count: the count is why some titles grew.
        p = cells[1].paragraphs[0]
        p.add_run(r["title"]).font.size = Pt(10)
        p2 = cells[1].add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        run2 = p2.add_run("%d characters, fits in 60" % r["tlen"])
        run2.font.size = Pt(7.5)
        run2.font.color.rgb = MUTED

        p = cells[2].paragraphs[0]
        p.add_run(r["desc"]).font.size = Pt(10)
        p2 = cells[2].add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        run2 = p2.add_run("%d characters, fits in 160" % r["dlen"])
        run2.font.size = Pt(7.5)
        run2.font.color.rgb = MUTED

        p = cells[3].paragraphs[0]
        if r["today"] == 0:
            run = p.add_run("No description at all")
        else:
            run = p.add_run("%d characters,\ncut off by Google" % r["today"])
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xA3, 0x3A, 0x1C)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
    return table


# ---------------------------------------------------------------- title block
rule()
h = doc.add_paragraph()
h.paragraph_format.space_after = Pt(4)
run = h.add_run("Search listings for empowerms.org")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = NAVY

body(
    "The short text that appears under each page's name in Google. "
    "Thirty-four addresses, for your approval.",
    size=12,
    color=MUTED,
    space_after=14,
)

body(
    "None of the sixteen main pages has one today. Google fills the gap with "
    "whatever text it can scrape off the page, which for these pages is "
    "usually the navigation. The eighteen staff and fellow biographies do have "
    "one, but each is the opening 350 to 420 characters of the biography, so "
    "Google cuts it off mid-sentence."
)
body(
    "Every line below is drawn from the page's own approved copy, or from the "
    "person's own biography. Nothing is invented, and no figure appears that is "
    "not already on the page it describes. Please mark up anything you want "
    "changed; anything unmarked we will take as approved."
)

# ------------------------------------------------------------------ the pages
rule()
heading("Part one: the sixteen site pages", space_before=4)
body(
    "Titles here are a little longer than the current ones. The current titles "
    "are short enough that Google leaves part of the result space unused.",
    color=MUTED,
    space_after=10,
)
build_table(DATA["pages"], show_today=True)

# ------------------------------------------------------------------- the bios
doc.add_page_break()
rule()
heading("Part two: the eighteen biographies", space_before=4)
body(
    "Each of these replaces a 350 to 420 character extract that Google "
    "truncates. The replacement leads with the person's role, then one line of "
    "background, taken from their own biography. The full biographies on the "
    "pages themselves are unchanged.",
    color=MUTED,
    space_after=10,
)
build_table(DATA["people"], show_today=True)

# -------------------------------------------------------------- the decisions
doc.add_page_break()
rule()
heading("Two decisions to confirm alongside the wording", space_before=4)

body(
    "1.  The home page is currently titled “Homepage (Elementor "
    "conversion)”. That is a label from the build process, and it is what a "
    "browser tab and a Google result both say today. The title in the table "
    "replaces it. Please confirm you are happy with the replacement wording."
)
body(
    "2.  Grant Callen's biography currently exists at two web addresses: "
    "%s/grant-callen/ and %s/person/grant-callen/. Both work, both are titled "
    "“Grant Callen”, and Google has no way to tell which is the real "
    "one, so the two compete with each other. We propose treating "
    "%s/person/grant-callen/ as the real one, which is the address the staff "
    "page already links to. The other stays reachable and simply points Google "
    "at it. No content is deleted either way." % (SITE, SITE, SITE)
)

rule()
heading("Not included, and why", size=14, space_before=4)
body(
    "Blog posts and the older pages are not in this list. They already carry "
    "descriptions, and while most are longer than Google will show, they are "
    "far better handled in one bulk pass than written out one by one. Happy to "
    "quote that separately if you want it done before launch."
)

body("Prepared 21 August 2026.", size=9, color=MUTED, italic=True)

out = HERE / "empower-search-listings.docx"
doc.save(out)
print("wrote", out)
